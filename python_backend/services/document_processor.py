import os
import tempfile
import asyncio
from typing import Dict, Any, List, Optional
from fastapi import UploadFile
import cv2
import numpy as np
from PIL import Image
import pytesseract
import easyocr
import pdfplumber
from docx import Document
import pandas as pd
import json

class DocumentProcessor:
    """
    Advanced document processor with OCR, computer vision, and AI capabilities
    """
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'],
            'text': ['.txt']
        }
        self.ocr_readers = {}
        self._initialize_ocr()
    
    def _initialize_ocr(self):
        """Initialize OCR engines"""
        try:
            # EasyOCR for better accuracy
            self.ocr_readers['easyocr'] = easyocr.Reader(['en'])
            # Tesseract for speed
            self.ocr_readers['tesseract'] = pytesseract
        except Exception as e:
            print(f"OCR initialization warning: {e}")
    
    async def process_document(
        self,
        file: UploadFile,
        analysis_type: str = "comprehensive",
        include_ocr: bool = True,
        include_layout: bool = True,
        include_tables: bool = True
    ) -> Dict[str, Any]:
        """
        Main document processing pipeline
        """
        try:
            # Determine file type
            file_extension = self._get_file_extension(file.filename)
            file_type = self._get_file_type(file_extension)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text based on file type
                if file_type == 'pdf':
                    extracted_data = await self._process_pdf(temp_file_path, include_tables)
                elif file_type == 'word':
                    extracted_data = await self._process_word(temp_file_path)
                elif file_type == 'image':
                    extracted_data = await self._process_image(temp_file_path, include_ocr)
                elif file_type == 'text':
                    extracted_data = await self._process_text(temp_file_path)
                else:
                    raise ValueError(f"Unsupported file type: {file_type}")
                
                # Add layout analysis if requested
                if include_layout and file_type in ['pdf', 'image']:
                    layout_data = await self._analyze_layout(temp_file_path)
                    extracted_data['layout_analysis'] = layout_data
                
                # Add metadata
                extracted_data['metadata'] = {
                    'filename': file.filename,
                    'file_type': file_type,
                    'file_size': len(content),
                    'analysis_type': analysis_type,
                    'processing_options': {
                        'ocr': include_ocr,
                        'layout': include_layout,
                        'tables': include_tables
                    }
                }
                
                return extracted_data
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            raise Exception(f"Document processing failed: {str(e)}")
    
    async def advanced_ocr(
        self,
        file: UploadFile,
        ocr_type: str = "auto",
        language: str = "en",
        extract_tables: bool = True,
        extract_charts: bool = True
    ) -> Dict[str, Any]:
        """
        Advanced OCR with multiple engines and table/chart extraction
        """
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Perform OCR based on type
                if ocr_type == "auto":
                    # Try multiple OCR engines and use the best result
                    ocr_results = {}
                    
                    # EasyOCR
                    try:
                        ocr_results['easyocr'] = await self._easyocr_extract(temp_file_path, language)
                    except:
                        pass
                    
                    # Tesseract
                    try:
                        ocr_results['tesseract'] = await self._tesseract_extract(temp_file_path, language)
                    except:
                        pass
                    
                    # Select best result based on confidence
                    best_result = self._select_best_ocr_result(ocr_results)
                    
                elif ocr_type == "easyocr":
                    best_result = await self._easyocr_extract(temp_file_path, language)
                elif ocr_type == "tesseract":
                    best_result = await self._tesseract_extract(temp_file_path, language)
                else:
                    raise ValueError(f"Unsupported OCR type: {ocr_type}")
                
                # Extract tables if requested
                if extract_tables:
                    tables = await self._extract_tables_from_image(temp_file_path)
                    best_result['extracted_tables'] = tables
                
                # Extract charts if requested
                if extract_charts:
                    charts = await self._extract_charts_from_image(temp_file_path)
                    best_result['extracted_charts'] = charts
                
                return best_result
                
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            raise Exception(f"Advanced OCR failed: {str(e)}")
    
    async def analyze_layout(
        self,
        file: UploadFile,
        detect_components: bool = True,
        analyze_structure: bool = True
    ) -> Dict[str, Any]:
        """
        Document layout analysis with computer vision
        """
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                layout_data = await self._analyze_layout(temp_file_path)
                
                if detect_components:
                    components = await self._detect_document_components(temp_file_path)
                    layout_data['components'] = components
                
                if analyze_structure:
                    structure = await self._analyze_document_structure(temp_file_path)
                    layout_data['structure'] = structure
                
                return layout_data
                
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            raise Exception(f"Layout analysis failed: {str(e)}")
    
    async def _process_pdf(self, file_path: str, include_tables: bool) -> Dict[str, Any]:
        """Process PDF documents"""
        try:
            with pdfplumber.open(file_path) as pdf:
                pages = []
                tables = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text() or ""
                    
                    # Extract tables if requested
                    page_tables = []
                    if include_tables:
                        page_tables = page.extract_tables()
                        for table in page_tables:
                            if table:
                                tables.append({
                                    'page': page_num + 1,
                                    'data': table,
                                    'rows': len(table),
                                    'columns': len(table[0]) if table else 0
                                })
                    
                    pages.append({
                        'page_number': page_num + 1,
                        'text': text,
                        'tables_count': len(page_tables),
                        'width': page.width,
                        'height': page.height
                    })
                
                return {
                    'document_type': 'pdf',
                    'total_pages': len(pages),
                    'pages': pages,
                    'extracted_tables': tables,
                    'total_tables': len(tables)
                }
                
        except Exception as e:
            raise Exception(f"PDF processing failed: {str(e)}")
    
    async def _process_word(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            doc = Document(file_path)
            
            paragraphs = []
            tables = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name,
                        'alignment': str(para.alignment)
                    })
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                tables.append({
                    'data': table_data,
                    'rows': len(table_data),
                    'columns': len(table_data[0]) if table_data else 0
                })
            
            return {
                'document_type': 'word',
                'paragraphs': paragraphs,
                'extracted_tables': tables,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables)
            }
            
        except Exception as e:
            raise Exception(f"Word document processing failed: {str(e)}")
    
    async def _process_image(self, file_path: str, include_ocr: bool) -> Dict[str, Any]:
        """Process image documents"""
        try:
            # Load image
            image = cv2.imread(file_path)
            if image is None:
                raise Exception("Failed to load image")
            
            # Basic image info
            height, width, channels = image.shape
            
            result = {
                'document_type': 'image',
                'image_info': {
                    'width': width,
                    'height': height,
                    'channels': channels,
                    'file_size': os.path.getsize(file_path)
                }
            }
            
            # Perform OCR if requested
            if include_ocr:
                ocr_result = await self._easyocr_extract(file_path, 'en')
                result['ocr_result'] = ocr_result
            
            return result
            
        except Exception as e:
            raise Exception(f"Image processing failed: {str(e)}")
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process text documents"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            words = content.split()
            
            return {
                'document_type': 'text',
                'content': content,
                'statistics': {
                    'characters': len(content),
                    'words': len(words),
                    'lines': len(lines),
                    'paragraphs': len([l for l in lines if l.strip()])
                }
            }
            
        except Exception as e:
            raise Exception(f"Text processing failed: {str(e)}")
    
    async def _easyocr_extract(self, file_path: str, language: str) -> Dict[str, Any]:
        """Extract text using EasyOCR"""
        try:
            reader = easyocr.Reader([language])
            results = reader.readtext(file_path)
            
            text_blocks = []
            full_text = ""
            
            for (bbox, text, confidence) in results:
                text_blocks.append({
                    'text': text,
                    'confidence': confidence,
                    'bbox': bbox
                })
                full_text += text + " "
            
            return {
                'ocr_engine': 'easyocr',
                'language': language,
                'text_blocks': text_blocks,
                'full_text': full_text.strip(),
                'total_blocks': len(text_blocks),
                'average_confidence': sum([b['confidence'] for b in text_blocks]) / len(text_blocks) if text_blocks else 0
            }
            
        except Exception as e:
            raise Exception(f"EasyOCR extraction failed: {str(e)}")
    
    async def _tesseract_extract(self, file_path: str, language: str) -> Dict[str, Any]:
        """Extract text using Tesseract"""
        try:
            # Set language
            pytesseract.pytesseract.tesseract_cmd = r'/usr/local/bin/tesseract'
            
            # Extract text
            text = pytesseract.image_to_string(file_path, lang=language)
            
            # Get confidence data
            data = pytesseract.image_to_data(file_path, lang=language, output_type=pytesseract.Output.DICT)
            
            # Calculate average confidence
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                'ocr_engine': 'tesseract',
                'language': language,
                'full_text': text,
                'average_confidence': avg_confidence,
                'total_words': len(text.split())
            }
            
        except Exception as e:
            raise Exception(f"Tesseract extraction failed: {str(e)}")
    
    async def _extract_tables_from_image(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract tables from image using computer vision"""
        try:
            # This is a simplified table extraction
            # In production, you'd use more sophisticated CV models
            image = cv2.imread(file_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Detect horizontal and vertical lines
            horizontal = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, -2)
            vertical = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, -2)
            
            # Combine horizontal and vertical lines
            table_structure = cv2.addWeighted(horizontal, 0.5, vertical, 0.5, 0.0)
            
            return [{
                'type': 'detected_table',
                'confidence': 0.7,  # Placeholder
                'bbox': [0, 0, image.shape[1], image.shape[0]]
            }]
            
        except Exception as e:
            return []
    
    async def _extract_charts_from_image(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract charts from image using computer vision"""
        try:
            # Simplified chart detection
            # In production, use specialized chart detection models
            image = cv2.imread(file_path)
            
            # Basic chart detection based on color and shape analysis
            # This is a placeholder - real implementation would use ML models
            
            return [{
                'type': 'potential_chart',
                'confidence': 0.6,
                'bbox': [0, 0, image.shape[1], image.shape[0]]
            }]
            
        except Exception as e:
            return []
    
    async def _analyze_layout(self, file_path: str) -> Dict[str, Any]:
        """Analyze document layout using computer vision"""
        try:
            image = cv2.imread(file_path)
            if image is None:
                return {}
            
            height, width, channels = image.shape
            
            # Basic layout analysis
            # In production, use LayoutLM or similar models
            
            # Detect text regions
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
            
            # Find contours
            contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Filter text regions
            text_regions = []
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                area = w * h
                if area > 100 and w > 10 and h > 10:  # Filter small regions
                    text_regions.append({
                        'x': x, 'y': y, 'width': w, 'height': h,
                        'area': area
                    })
            
            return {
                'layout_type': 'detected',
                'image_dimensions': {'width': width, 'height': height},
                'text_regions': text_regions,
                'total_regions': len(text_regions)
            }
            
        except Exception as e:
            return {}
    
    async def _detect_document_components(self, file_path: str) -> List[Dict[str, Any]]:
        """Detect document components (headers, footers, tables, etc.)"""
        # Placeholder for component detection
        return [
            {'type': 'header', 'confidence': 0.8, 'bbox': [0, 0, 100, 50]},
            {'type': 'body', 'confidence': 0.9, 'bbox': [0, 50, 100, 200]},
            {'type': 'footer', 'confidence': 0.7, 'bbox': [0, 200, 100, 250]}
        ]
    
    async def _analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze overall document structure"""
        # Placeholder for structure analysis
        return {
            'structure_type': 'standard',
            'sections': ['header', 'body', 'footer'],
            'layout_complexity': 'medium'
        }
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        if not filename:
            return '.txt'
        return os.path.splitext(filename)[1].lower()
    
    def _get_file_type(self, extension: str) -> str:
        """Get file type from extension"""
        for file_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return file_type
        return 'text'  # Default to text
    
    def _select_best_ocr_result(self, ocr_results: Dict[str, Any]) -> Dict[str, Any]:
        """Select the best OCR result based on confidence"""
        if not ocr_results:
            return {}
        
        best_result = None
        best_confidence = 0
        
        for engine, result in ocr_results.items():
            if 'average_confidence' in result and result['average_confidence'] > best_confidence:
                best_confidence = result['average_confidence']
                best_result = result
        
        return best_result or list(ocr_results.values())[0]
